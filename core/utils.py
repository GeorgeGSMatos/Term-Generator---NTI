"""
Módulo de Funcionalidades e Utilitários (Core Utilities).

Fornece manipuladores do sistema de arquivos, conversores de datas, motores
de automação híbrida Word/PDF e formatadores agnósticos desacoplados da UI.
"""

# ==============================================================================
# 1. IMPORTS E DEPENDÊNCIAS
# ==============================================================================

import os
import re
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional, Tuple

# --- 1.1. Dependências de Automação Windows COM+ ---
try:
    import comtypes.client
    import pythoncom
except ImportError:
    comtypes = None
    pythoncom = None

from docx import Document
from docxtpl import DocxTemplate

# --- 1.2. Configurações Core ---
from core.settings import MONTH_MAP, NETWORK_FOLDERS, OP_STYLES, load_setting

# ==============================================================================
# 2. CONTROLADORES DE FLUXO E NEGÓCIO
# ==============================================================================


def fire_progress(
    callback: Optional[Callable[[int, str], None]], pct: int, msg: str
) -> None:
    """Dispara o callback de progresso de forma segura.

    Args:
        callback (Optional[Callable]): Função de callback da UI.
        pct (int): Percentual de conclusão (0-100).
        msg (str): Mensagem descritiva do progresso.
    """
    if callback:
        callback(pct, msg)


def classify_operation_type(op_val: str, obs_val: str, keywords: List[str]) -> str:
    """Classifica o tipo final de operação baseado em contexto textual.

    Detecta se a operação deve ser classificada como 'Movimentação' baseando-se
    em palavras-chave encontradas nas observações.

    Args:
        op_val (str): Tipo de operação original da UI.
        obs_val (str): Observações digitadas pelo usuário.
        keywords (List[str]): Lista de palavras-chave indicativas.

    Returns:
        str: Tipo de operação classificado.
    """
    is_movement = any(kw in obs_val.lower() for kw in keywords)
    return "Movimentação" if is_movement else op_val


# ==============================================================================
# 3. FORMATADORES E CONVERSORES
# ==============================================================================


def format_display_date(date_str: str) -> str:
    """Converte uma data ISO (YYYY-MM-DD) para formato legível (DD/MM/YYYY).

    Args:
        date_str (str): Data em formato ISO ou string do banco.

    Returns:
        str: Data formatada ou '--/--/----' se inválida.
    """
    try:
        if date_str:
            dt_obj = datetime.strptime(str(date_str).split()[0], "%Y-%m-%d")
            return dt_obj.strftime("%d/%m/%Y")
    except (ValueError, IndexError, AttributeError):
        pass
    return "--/--/----"


def standardize_iso_date(text_or_date: Optional[str]) -> str:
    """Identifica e extrai a ÚLTIMA data válida em um texto para YYYY-MM-DD.

    Suporta formatos por extenso (Ex: 01 de Janeiro de 2024), numéricos
    com vários delimitadores e anos abreviados.

    Args:
        text_or_date (Optional[str]): Texto bruto para busca.

    Returns:
        str: Data ISO formatada ou string vazia.
    """
    if not text_or_date or str(text_or_date).strip() in [
        "",
        "S/N",
        "null",
        "None",
    ]:
        return ""

    raw_text = str(text_or_date).strip()
    mapa_meses = {
        "jan": 1,
        "fev": 2,
        "mar": 3,
        "abr": 4,
        "mai": 5,
        "jun": 6,
        "jul": 7,
        "ago": 8,
        "set": 9,
        "out": 10,
        "nov": 11,
        "dez": 12,
    }

    # --- 3.1. Heurística: Datas por Extenso ---
    matches_ext = re.findall(
        r"(\d{1,2})\s*(?:de)?\s*([a-zA-ZçÇ]+)\s*(?:de)?\s+(\d{4})",
        raw_text,
        re.IGNORECASE,
    )
    if matches_ext:
        dia, mes_str, ano = (
            int(matches_ext[-1][0]),
            matches_ext[-1][1].lower(),
            int(matches_ext[-1][2]),
        )
        mes_num = next((num for nome, num in mapa_meses.items() if nome in mes_str), 1)
        if 2010 <= ano <= 2035:
            return f"{ano:04d}-{mes_num:02d}-{dia:02d}"

    # --- 3.2. Heurística: Datas Numéricas ---
    matches_num = re.findall(r"(\d{1,2})[\/\-\.](\d{1,2})[\/\-\.](\d{2,4})", raw_text)
    if matches_num:
        p1, p2, p3 = matches_num[-1]
        dia, mes, ano = int(p1), int(p2), int(p3)
        if len(str(p3)) == 2:
            ano += 2000

        if 1 <= mes <= 12 and 1 <= dia <= 31 and 2010 <= ano <= 2035:
            return f"{ano:04d}-{mes:02d}-{dia:02d}"

    return ""


def format_ticker_time(time_raw: str) -> Tuple[str, str]:
    """Extrai Hora e Dia de um Datetime para o Dashboard.

    Args:
        time_raw (str): String no formato %Y-%m-%d %H:%M:%S.

    Returns:
        Tuple[str, str]: Par (Hora, Dia) formatados.
    """
    try:
        dt_parsed = datetime.strptime(time_raw, "%Y-%m-%d %H:%M:%S")
        h_fmt = dt_parsed.strftime("%H:%M")
        d_fmt = dt_parsed.strftime("%d/%m")
        return h_fmt, d_fmt
    except ValueError:
        return "--:--", str(time_raw)[:10]


def clean_filename(filename: str) -> str:
    """Remove extensões redundantes de nomes de arquivos para a UI.

    Args:
        filename (str): Nome bruto do arquivo.

    Returns:
        str: Nome limpo sem extensões comuns (.docx, .pdf).
    """
    display_text = str(filename)
    for ext in [".docx", ".pdf", ".doc"]:
        if display_text.lower().endswith(ext):
            return display_text[: -len(ext)]
    return display_text


def resolve_operation_context(
    op_type: str, doc_path: str, obs_text: str
) -> Tuple[str, str, bool]:
    """Infere o contexto visual (ícone/cor) do card baseado na operação.

    Args:
        op_type (str): Categoria da operação.
        doc_path (str): Caminho do arquivo para análise de pasta.
        obs_text (str): Observações para análise semântica.

    Returns:
        Tuple[str, str, bool]: (Ícone, Cor, Flag de Lógica Híbrida).
    """
    op_type_str = str(op_type)

    if op_type_str != "Movimentação":
        style = OP_STYLES.get(op_type_str, OP_STYLES.get("Desconhecido"))
        return style["icon"], style["color"], False

    path_lower = str(doc_path or "").lower()
    obs_lower = str(obs_text or "").lower()

    is_return_context = (
        "devolu" in path_lower or "td -" in path_lower or "devolvido" in obs_lower
    )

    if is_return_context:
        style_ctx = OP_STYLES.get("Devolução")
        return style_ctx["icon"], "red", True
    else:
        style_ctx = OP_STYLES.get("Entrega")
        return style_ctx["icon"], style_ctx["color"], True


# ==============================================================================
# 4. MANIPULAÇÃO DE ARQUIVOS
# ==============================================================================


def sanitize_filename(filename: str) -> str:
    """Esteriliza nomes de arquivos contra caracteres proibidos no Windows.

    Args:
        filename (str): Sugestão de nome para o arquivo.

    Returns:
        str: Nome sanitizado.
    """
    if not filename:
        return "SemTitulo"

    return re.sub(r'[\\/*?:"<>|]', "", str(filename)).strip()


def get_full_date_text() -> str:
    """Gera o rodapé temporal descritivo (Ex: Salvador, 22 de Março de 2024).

    Returns:
        str: Data por extenso localizada.
    """
    now: datetime = datetime.now()
    setting = load_setting()
    city = setting.get("cidade", "Salvador")
    month_name = MONTH_MAP.get(now.month, "Janeiro")

    return f"{city}, {now.day} de {month_name} de {now.year}"


def calculate_smart_path(root_path: str, operation_type: str) -> str:
    """Calcula dinamicamente a estrutura de pastas baseada na operação e data.

    Args:
        root_path (str): Pasta raiz configurada na rede.
        operation_type (str): Tipo de operação (Entrega, Devolução, etc).

    Returns:
        str: Caminho absoluto validado para o destino.
    """
    now = datetime.now()
    year_str = str(now.year)
    month_name = MONTH_MAP.get(now.month, "Janeiro")

    folder_info = NETWORK_FOLDERS.get(
        operation_type, {"pasta": "OUTROS", "prefixo": "DOC"}
    )
    folder_main = folder_info["pasta"]
    prefix = folder_info["prefixo"]

    path = os.path.join(root_path, folder_main)
    annual_folder = f"TERMO DE {operation_type.upper()} {year_str}"
    path = os.path.join(path, annual_folder)

    monthly_folder = f"{prefix} - {month_name}"
    path = os.path.join(path, monthly_folder)

    if not os.path.exists(path):
        try:
            os.makedirs(path, exist_ok=True)
        except OSError as e:
            print(f"⚠️ Erro ao criar pastas em {path}: {e}")
            return root_path

    return path


def truncate_to_max_path(base_folder: str, filename: str) -> str:
    """Respeita o limite MAX_PATH do Windows (260 caracteres).

    Args:
        base_folder (str): Pasta de destino.
        filename (str): Nome do arquivo pretendido.

    Returns:
        str: Nome truncado se necessário para garantir segurança no FS.
    """
    MAX_SAFE_LENGTH = 250
    full_path = os.path.join(base_folder, filename)

    if len(full_path) <= MAX_SAFE_LENGTH:
        return filename

    name, ext = os.path.splitext(filename)
    excess_chars = len(full_path) - MAX_SAFE_LENGTH
    new_length = max(10, len(name) - excess_chars)
    truncated_name = name[:new_length].strip()

    return f"{truncated_name}{ext}"


def generate_unique_path(
    folder: str, base_name: str, extension: str
) -> Tuple[str, str]:
    """Gera um caminho único evitando sobrescrita via versionamento incremental.

    Args:
        folder (str): Pasta de destino.
        base_name (str): Nome base do arquivo.
        extension (str): Extensão (ex: '.docx').

    Returns:
        Tuple[str, str]: (Caminho Absoluto, Nome Único).
    """
    counter: int = 1
    clean_name: str = sanitize_filename(base_name)
    filename: str = f"{clean_name}{extension}"
    full_path: str = os.path.join(folder, filename)

    while os.path.exists(full_path):
        filename = f"{clean_name} ({counter}){extension}"
        full_path = os.path.join(folder, filename)
        counter += 1

    return full_path, filename


def check_word_availability() -> bool:
    """Verifica se as bibliotecas COM+ do Windows estão presentes.

    Returns:
        bool: True se disponível, False caso contrário.
    """
    return comtypes is not None


# ==============================================================================
# 5. MOTORES DE GERAÇÃO E EXTRAÇÃO DOCUMENTAL
# ==============================================================================


def extract_docx_text(doc_path: str) -> str:
    """Extrai texto bruto de documentos Word (.docx) via OpenXML.

    Varre parágrafos, tabelas, rodapés e caixas de texto internas.

    Args:
        doc_path (str): Caminho absoluto para o arquivo .docx.

    Returns:
        str: Texto consolidado do documento ou string vazia em caso de erro.
    """
    try:
        if not os.path.exists(doc_path):
            return ""

        doc = Document(doc_path)
        full_text: List[str] = []

        # --- 5.1. Parágrafos ---
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())

        # --- 5.2. Tabelas ---
        for table in doc.tables:
            for row in table.rows:
                row_data = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_data:
                    full_text.append(" | ".join(row_data))

        # --- 5.3. Rodapés e Metadados ---
        for section in doc.sections:
            footers = [
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            for footer in footers:
                if footer:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            full_text.append(p.text.strip())

        # --- 5.4. Caixas de Texto (txbxContent) ---
        for child in doc.element.body.iter():
            if child.tag.endswith("txbxContent"):
                for p in child.iter():
                    if p.tag.endswith("p"):
                        p_text = "".join(
                            [t.text for t in p.iter() if t.tag.endswith("t") and t.text]
                        )
                        if p_text.strip():
                            full_text.append(p_text.strip())

        return "\n".join(full_text)
    except Exception as e:
        print(f"⚠️ Erro ao extrair texto de '{os.path.basename(doc_path)}': {e}")
        return ""


def generate_final_docx(
    context: Dict[str, Any], template_path: str, output_path: str
) -> bool:
    """Renderiza dados em um Template Word via Jinja2 (DocxTemplate).

    Args:
        context (Dict[str, Any]): Contexto de dados para substituição.
        template_path (str): Caminho do arquivo .docx original.
        output_path (str): Destino final do arquivo preenchido.

    Returns:
        bool: True se gerado com sucesso.
    """
    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"❌ Erro Crítico no Motor Word: {e}")
        raise e


def convert_to_pdf(input_docx: str, output_pdf: str) -> bool:
    """Converte um DOCX para PDF utilizando Automação Word COM em modo Headless.

    Garante o encerramento do processo Winword e a liberação de recursos.

    Args:
        input_docx (str): Caminho do arquivo Word de origem.
        output_pdf (str): Caminho do arquivo PDF de destino.

    Returns:
        bool: True se a conversão foi concluída com sucesso.
    """
    if not os.path.exists(input_docx):
        print(f"❌ Motor PDF: Arquivo de origem ausente ({input_docx})")
        return False

    if not comtypes or not pythoncom:
        print("⚠️ Bibliotecas COM ausentes. Conversão PDF indisponível.")
        return False

    word_app = None
    doc_obj = None

    try:
        pythoncom.CoInitialize()

        word_app = comtypes.client.CreateObject("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0

        input_abs: str = os.path.abspath(input_docx)
        output_abs: str = os.path.abspath(output_pdf)

        doc_obj = word_app.Documents.Open(input_abs)
        doc_obj.ExportAsFixedFormat(
            OutputFileName=output_abs,
            ExportFormat=17,  # constant wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
            DocStructureTags=True,
        )

        return True

    except Exception as e:
        print(f"❌ Erro no Motor PDF: {e}")
        return False

    finally:
        if doc_obj:
            try:
                doc_obj.Close(SaveChanges=0)
            except Exception:
                pass

        if word_app:
            try:
                word_app.Quit()
            except Exception:
                pass

        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
